from __future__ import annotations

import io
import os
from typing import Any

import pandas as pd
import streamlit as st
from openpyxl.styles import Font
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from core.free_llm_client import FreeLLMClient
from core.llm_task_parser import parse_task_prompt_llm_first
from core.models import ProviderSettings
from pipeline.orchestrator import SearchOrchestrator

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None


# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Research Agent",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Global CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.stApp { background: linear-gradient(180deg, #070b14 0%, #0b1020 100%); color: #e7ecf5; }
.block-container { padding-top: 4.8rem; padding-bottom: 2rem; max-width: 1180px; }

/* Prevent Streamlit's floating top toolbar / sidebar toggle from sitting on top of the hero title */
header[data-testid="stHeader"] {
    background: rgba(7,11,20,0.98) !important;
    height: 3.4rem !important;
    border-bottom: 1px solid rgba(124,143,179,0.08);
}
[data-testid="stToolbar"] {
    right: 0.75rem !important;
}
[data-testid="collapsedControl"] {
    top: 0.85rem !important;
    left: 0.85rem !important;
    color: #9cc2ff !important;
    z-index: 999999 !important;
}
button[kind="header"] { color: #9cc2ff !important; }

@media (max-width: 900px) {
    .block-container { padding-top: 5.2rem; }
    .hero-title { font-size: 1.75rem; }
}

/* ── Hero ── */
.hero-wrap { display: flex; align-items: flex-start; justify-content: space-between;
             margin-bottom: 0.7rem; flex-wrap: wrap; gap: 0.4rem; }
.hero-title { font-size: 2.1rem; font-weight: 800; letter-spacing: -0.03em; color: #f8fafc; margin: 0; }
.hero-subtitle { color: #8fa5c8; font-size: 0.97rem; margin: 0.18rem 0 0 0; line-height: 1.4; }
.hero-badge { display: inline-flex; align-items: center; background: rgba(83,113,255,0.12);
              border: 1px solid rgba(83,113,255,0.22); border-radius: 999px;
              padding: 0.22rem 0.75rem; font-size: 0.77rem; color: #a8c0ff; margin-top: 0.15rem; white-space: nowrap; }

/* ── Capability tiers card ── */
.tier-card {
    background: linear-gradient(135deg, rgba(20,30,55,0.95) 0%, rgba(12,18,35,0.95) 100%);
    border: 1px solid rgba(124,143,179,0.18); border-radius: 18px;
    padding: 1rem 1.15rem 0.9rem 1.15rem; margin-bottom: 1rem;
}
.tier-card-header { display: flex; align-items: flex-start; justify-content: space-between;
                    gap: 0.75rem; margin-bottom: 0.8rem; flex-wrap: wrap; }
.tier-card-title { color: #f0f5ff; font-weight: 700; font-size: 0.97rem; margin: 0; }
.tier-card-sub   { color: #7a8fa8; font-size: 0.82rem; margin-top: 0.18rem; }
.sidebar-cta {
    display: inline-flex; align-items: center; gap: 0.35rem; white-space: nowrap;
    background: rgba(83,113,255,0.16); border: 1px solid rgba(83,113,255,0.28);
    color: #b8caff; border-radius: 10px; padding: 0.3rem 0.7rem;
    font-size: 0.8rem; font-weight: 600; flex-shrink: 0;
}
.tiers-grid { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 0.6rem; }
.tier-cell { background: rgba(255,255,255,0.04); border: 1px solid rgba(124,143,179,0.13);
             border-radius: 12px; padding: 0.65rem 0.8rem; }
.tier-cell-label { font-size: 0.71rem; font-weight: 600; letter-spacing: 0.05em;
                   text-transform: uppercase; margin-bottom: 0.22rem; }
.tier-cell-name  { font-size: 0.88rem; font-weight: 700; margin-bottom: 0.18rem; }
.tier-cell-desc  { font-size: 0.76rem; color: #7a8fa8; line-height: 1.35; }
.tier-free { border-color: rgba(100,180,130,0.22); }
.tier-llm  { border-color: rgba(130,100,220,0.28); }
.tier-full { border-color: rgba(83,113,255,0.3); }
.tier-free .tier-cell-label { color: #6db88a; }
.tier-llm  .tier-cell-label { color: #a07ee0; }
.tier-full .tier-cell-label { color: #6e9dff; }
.tier-free .tier-cell-name  { color: #a8dbbe; }
.tier-llm  .tier-cell-name  { color: #c5aaff; }
.tier-full .tier-cell-name  { color: #98bdff; }

/* ── Search shell ── */
.search-shell {
    background: rgba(15,22,38,0.92); border: 1px solid rgba(124,143,179,0.18);
    border-radius: 20px; padding: 1rem 1rem 0.95rem 1rem;
    box-shadow: 0 18px 44px rgba(0,0,0,0.28); margin-bottom: 0.85rem;
}

/* ── Quick-fill label ── */
.chips-label { color: #5e7a9e; font-size: 0.73rem; font-weight: 600; margin-bottom: 0.32rem;
               letter-spacing: 0.04em; text-transform: uppercase; }

/* ── History strip ── */
.history-label { color: #4a6078; font-size: 0.73rem; font-weight: 600;
                 text-transform: uppercase; letter-spacing: 0.04em; margin-bottom: 0.35rem; }

/* ── Summary card ── */
.summary-card { background: rgba(15,22,38,0.9); border: 1px solid rgba(124,143,179,0.16);
                border-radius: 16px; padding: 0.85rem 1.1rem 0.75rem 1.1rem; margin-bottom: 0.85rem; }
.summary-card-title { color: #f8fafc; font-weight: 700; font-size: 0.97rem; margin-bottom: 0.55rem; }
.chips-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 0.35rem 0.45rem; margin-bottom: 0.45rem; }
.metric-chip { display: flex; align-items: center; padding: 0.27rem 0.65rem; border-radius: 999px;
               background: rgba(83,113,255,0.13); border: 1px solid rgba(83,113,255,0.2);
               color: #dbe6ff; font-size: 0.8rem; overflow: hidden; }
.chip-label { color: #8fa5d0; font-weight: 500; margin-right: 0.3em; flex-shrink: 0; }
.chip-value { color: #e2ecff; font-weight: 600; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.summary-requested { color: #5e7a9e; font-size: 0.78rem; margin-top: 0.2rem; }

/* ── Results status bar ── */
.results-bar { display: flex; align-items: stretch; background: rgba(15,22,38,0.9);
               border: 1px solid rgba(124,143,179,0.15); border-radius: 14px;
               padding: 0.65rem 1rem; margin-bottom: 0.85rem; flex-wrap: wrap; gap: 0.5rem; }
.results-bar-item { display: flex; flex-direction: column; padding: 0 1.1rem;
                    border-right: 1px solid rgba(124,143,179,0.1); }
.results-bar-item:first-child { padding-left: 0; }
.results-bar-item:last-child  { border-right: none; }
.rb-label { color: #5e7a9e; font-size: 0.68rem; font-weight: 600;
            text-transform: uppercase; letter-spacing: 0.05em; }
.rb-value { color: #f0f5ff; font-size: 1.12rem; font-weight: 700; line-height: 1.2; }
.rb-sub   { color: #7a8fa8; font-size: 0.73rem; }

/* ── Provider status dots ── */
.provider-row { display: flex; flex-wrap: wrap; gap: 0.3rem; margin-top: 0.55rem; }
.provider-dot { display: inline-flex; align-items: center; gap: 0.3rem; font-size: 0.78rem;
                padding: 0.2rem 0.5rem; border-radius: 8px; border: 1px solid; }
.provider-dot.on  { color: #a8dbbe; background: rgba(41,89,63,0.18); border-color: rgba(91,166,116,0.25); }
.provider-dot.off { color: #3a4d62; background: rgba(255,255,255,0.03); border-color: rgba(124,143,179,0.09); }
.dot-circle { width: 6px; height: 6px; border-radius: 50%; display: inline-block; flex-shrink: 0; }
.dot-circle.on  { background: #5cb885; box-shadow: 0 0 5px rgba(92,184,133,0.55); }
.dot-circle.off { background: #2a3b4e; }

/* ── Key guide inside expander ── */
.key-guide { background: rgba(10,16,30,0.88); border-radius: 12px;
             padding: 0.8rem 0.9rem; margin-bottom: 0.8rem;
             border: 1px solid rgba(124,143,179,0.12); }
.kgp { margin-bottom: 0.75rem; padding-bottom: 0.65rem; border-bottom: 1px solid rgba(124,143,179,0.08); }
.kgp:last-child { margin-bottom: 0; padding-bottom: 0; border-bottom: none; }
.kg-name  { font-size: 0.84rem; font-weight: 700; color: #c5d4ff; margin-bottom: 0.1rem; }
.kg-role  { font-size: 0.74rem; color: #6e8db5; margin-bottom: 0.25rem; }
.kg-steps { list-style: none; padding: 0; margin: 0; }
.kg-steps li { font-size: 0.76rem; color: #8fa5c8; padding: 0.08rem 0;
               display: flex; align-items: flex-start; gap: 0.4rem; }
.kg-steps li .sn { color: #4a6882; font-weight: 700; flex-shrink: 0; min-width: 1rem; }
.kg-free { display: inline-block; background: rgba(41,89,63,0.28); color: #7ec89a;
           border-radius: 6px; font-size: 0.66rem; font-weight: 600;
           padding: 0.06rem 0.32rem; margin-left: 0.3rem; vertical-align: middle; }
.key-guide a {
    color: #9cc2ff;
    font-weight: 700;
    text-decoration: none;
    border-bottom: 1px dashed rgba(156,194,255,0.55);
}
.key-guide a:hover {
    color: #ffffff;
    border-bottom-color: #ffffff;
}

/* ── Key status ── */
.key-status { color: #d7eadc; background: rgba(41,89,63,0.22); border: 1px solid rgba(91,166,116,0.18);
              border-radius: 12px; padding: 0.5rem 0.75rem; font-size: 0.84rem; margin-top: 0.5rem; }

/* ── Empty state ── */
.empty-box { background: rgba(15,22,38,0.88); border: 1px dashed rgba(124,143,179,0.22);
             border-radius: 16px; padding: 1.5rem 1.25rem; }
.empty-icon  { font-size: 2rem; margin-bottom: 0.45rem; }
.empty-title { color: #d0dcee; font-weight: 700; font-size: 1rem; margin-bottom: 0.35rem; }
.empty-desc  { color: #7a8fa8; font-size: 0.87rem; line-height: 1.5; }
.empty-tips  { margin-top: 0.7rem; }
.empty-tip   { display: flex; align-items: flex-start; gap: 0.5rem; color: #8fa5c8;
               font-size: 0.83rem; padding: 0.22rem 0; }
.tip-icon    { flex-shrink: 0; }

/* ── Download box ── */
.download-box { background: rgba(15,22,38,0.9); border: 1px solid rgba(124,143,179,0.16);
                border-radius: 16px; padding: 1rem; margin-top: 1rem; }

/* ── Sidebar footer ── */
.sidebar-footer { color: #2e3f52; font-size: 0.72rem; text-align: center; margin-top: 1rem;
                  line-height: 1.55; border-top: 1px solid rgba(124,143,179,0.07); padding-top: 0.8rem; }

/* ── Misc ── */
.stButton > button, .stDownloadButton > button { border-radius: 12px; }
</style>
""", unsafe_allow_html=True)


# ╔══════════════════════════════════════════════════════════╗
# ║  PURE HELPERS                                            ║
# ╚══════════════════════════════════════════════════════════╝

def _secret(key: str, default: str = "") -> str:
    try:
        return st.secrets.get(key, "") or os.getenv(key, default)
    except Exception:
        return os.getenv(key, default)


def _clean(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    return "" if s.lower() in {"nan", "none", "null"} else s


def _normalize_filename(name: str, ext: str) -> str:
    base = (name or "results").strip() or "results"
    for e in [".xlsx", ".csv", ".json", ".pdf", ".docx", ".doc"]:
        if base.lower().endswith(e):
            base = base[: -len(e)]
            break
    return f"{base}{ext}"


def _human_task(task_type: str) -> str:
    return {
        "entity_discovery":        "Companies",
        "document_research":       "Academic Papers",
        "people_search":           "LinkedIn Profiles",
        "market_research":         "Market Research",
        "entity_enrichment":       "Enrichment",
        "similar_entity_expansion":"Similar Entities",
    }.get(task_type or "", "Results")


def _human_category(cat: str) -> str:
    return {
        "software_company": "Digital / software companies",
        "service_company":  "Service / engineering companies",
        "general":          "General results",
    }.get((cat or "general").strip(), (cat or "general").replace("_", " ").title())


def _read_file(file) -> pd.DataFrame | None:
    if file is None:
        return None
    try:
        if file.name.lower().endswith(".csv"):
            return pd.read_csv(file)
        return pd.read_excel(file)
    except Exception:
        return None


def _connected_integrations(keys: dict[str, str]) -> list[str]:
    labels = {"groq": "Groq", "gemini": "Gemini", "exa": "Exa", "tavily": "Tavily", "serpapi": "SerpApi"}
    return [labels[k] for k, v in keys.items() if str(v).strip()]


# ── Render helpers ────────────────────────────────────────────────────────────

def _chip(label: str, value: str) -> str:
    return (
        f'<div class="metric-chip">'
        f'<span class="chip-label">{label}:</span>'
        f'<span class="chip-value">{value}</span>'
        f'</div>'
    )


def _render_search_summary(task_spec: dict) -> None:
    geo  = task_spec.get("geography", {}) or {}
    inc  = [str(x).title() for x in (geo.get("include_countries")          or []) if str(x).strip()]
    exc  = [str(x).title() for x in (geo.get("exclude_countries")          or []) if str(x).strip()]
    excp = [str(x).title() for x in (geo.get("exclude_presence_countries") or []) if str(x).strip()]

    pairs: list[tuple[str, str]] = [
        ("Type",     _human_task(task_spec.get("task_type", ""))),
        ("Category", _human_category(task_spec.get("target_category", ""))),
    ]
    if _clean(task_spec.get("industry")):
        pairs.append(("Focus", _clean(task_spec.get("industry"))))
    if inc:
        pairs.append(("Region", ", ".join(inc)))
    if exc:
        pairs.append(("Excl. HQ", ", ".join(exc)))
    if excp:
        pairs.append(("Excl. presence", ", ".join(excp)))

    if len(pairs) % 2 != 0:
        pairs.append(("", ""))

    chips_html = "".join(
        _chip(lbl, val) if lbl else '<div></div>'
        for lbl, val in pairs
    )
    requested = [str(x) for x in (task_spec.get("target_attributes") or []) if str(x).strip()]
    req_html  = (
        f'<div class="summary-requested">📋 Requested fields: {", ".join(requested)}</div>'
        if requested else ""
    )
    st.markdown(
        f'<div class="summary-card">'
        f'<div class="summary-card-title">🔍 Search interpreted as</div>'
        f'<div class="chips-grid">{chips_html}</div>'
        f'{req_html}'
        f'</div>',
        unsafe_allow_html=True,
    )


def _render_results_bar(n_results: int, mode: str, coverage: int, task_type: str) -> None:
    st.markdown(
        f'<div class="results-bar">'
        f'<div class="results-bar-item"><span class="rb-label">Results found</span>'
        f'<span class="rb-value">{n_results}</span></div>'
        f'<div class="results-bar-item"><span class="rb-label">Search mode</span>'
        f'<span class="rb-value">{mode}</span></div>'
        f'<div class="results-bar-item"><span class="rb-label">Coverage</span>'
        f'<span class="rb-value">{coverage}</span><span class="rb-sub">candidates scanned</span></div>'
        f'<div class="results-bar-item"><span class="rb-label">Type</span>'
        f'<span class="rb-value">{_human_task(task_type)}</span></div>'
        f'</div>',
        unsafe_allow_html=True,
    )


def _render_empty_state(task_type: str) -> None:
    tips_map = {
        "entity_discovery": [
            ("🌐", "Add <strong>Exa</strong> or <strong>Tavily</strong> in the sidebar for broader web coverage"),
            ("✍️", "Try rephrasing — e.g. add 'service providers', 'contractors', or 'suppliers'"),
            ("🔭", "Switch to <strong>Balanced</strong> or <strong>Deep</strong> mode and raise coverage"),
        ],
        "document_research": [
            ("📡", "<strong>Exa</strong> is the strongest source for academic papers — add a key in the sidebar"),
            ("🔎", "Include 'SPE paper', 'abstract', or a specific journal in your query"),
            ("🔭", "Switch to <strong>Deep</strong> mode for hard-to-find literature"),
        ],
        "people_search": [
            ("🔗", "LinkedIn search works best with <strong>Exa</strong> — add a key in the sidebar"),
            ("✍️", "Be specific: job title + company type + country"),
            ("🔭", "Try <strong>Balanced</strong> or <strong>Deep</strong> mode"),
        ],
    }
    tips = tips_map.get(task_type, [
        ("🌐", "Add Exa or Tavily in the sidebar for broader coverage"),
        ("✍️", "Try rephrasing with more specific terms"),
        ("🔭", "Increase coverage or switch to Deep mode"),
    ])
    tips_html = "".join(
        f'<div class="empty-tip"><span class="tip-icon">{icon}</span><span>{text}</span></div>'
        for icon, text in tips
    )
    st.markdown(
        f'<div class="empty-box">'
        f'<div class="empty-icon">🔍</div>'
        f'<div class="empty-title">No strong matches found</div>'
        f'<div class="empty-desc">The search completed but no results cleared the relevance threshold.</div>'
        f'<div class="empty-tips">{tips_html}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


def _summary_from_record(r: dict) -> str:
    for key in ["summary", "description", "snippet", "notes", "abstract"]:
        val = _clean(r.get(key))
        if val:
            return val
    return ""


def _build_display_df(records: list[dict], task_meta: dict) -> pd.DataFrame:
    task_type   = task_meta.get("task_type", "")
    entity_type = (task_meta.get("target_entity_types") or [""])[0]
    rows = []
    for r in records:
        if task_type == "document_research" or entity_type == "paper":
            rows.append({
                "Title":   _clean(r.get("company_name")) or _clean(r.get("title")),
                "Link":    _clean(r.get("website")) or _clean(r.get("source_url")),
                "Authors": _clean(r.get("authors")),
                "Year":    _clean(r.get("publication_year")),
                "Summary": _summary_from_record(r)[:320],
            })
        elif task_type == "people_search" or entity_type == "person":
            rows.append({
                "Name":      _clean(r.get("company_name")) or _clean(r.get("title")),
                "Link":      _clean(r.get("linkedin_profile")) or _clean(r.get("linkedin_url")) or _clean(r.get("website")),
                "Employer":  _clean(r.get("employer_name")),
                "Job title": _clean(r.get("job_title")),
                "Location":  _clean(r.get("city")) or _clean(r.get("country")),
                "Summary":   _summary_from_record(r)[:260],
            })
        else:
            rows.append({
                "Name":     _clean(r.get("company_name")) or _clean(r.get("title")),
                "Link":     _clean(r.get("website")) or _clean(r.get("source_url")),
                "Email":    _clean(r.get("email")),
                "Phone":    _clean(r.get("phone")),
                "LinkedIn": _clean(r.get("linkedin_profile")) or _clean(r.get("linkedin_url")),
                "Summary":  _summary_from_record(r)[:240],
            })
    return pd.DataFrame(rows)


def _build_export_df(records: list[dict], task_meta: dict) -> pd.DataFrame:
    rows = []
    task_type   = task_meta.get("task_type", "")
    entity_type = (task_meta.get("target_entity_types") or [""])[0]
    for r in records:
        link     = _clean(r.get("website")) or _clean(r.get("source_url"))
        linkedin = _clean(r.get("linkedin_profile")) or _clean(r.get("linkedin_url"))
        if task_type == "people_search" or entity_type == "person":
            link = linkedin or link
        rows.append({
            "name":     _clean(r.get("company_name")) or _clean(r.get("title")),
            "link":     link,
            "email":    _clean(r.get("email")),
            "phone":    _clean(r.get("phone")),
            "linkedin": linkedin,
            "summary":  _summary_from_record(r),
        })
    return pd.DataFrame(rows)


def _build_paper_summaries_df(records: list[dict]) -> pd.DataFrame:
    rows = []
    for r in records:
        rows.append({
            "Title": _clean(r.get("company_name")) or _clean(r.get("title")) or "Untitled",
            "Link": _clean(r.get("website")) or _clean(r.get("source_url")),
            "Authors": _clean(r.get("authors")),
            "Year": _clean(r.get("publication_year")),
            "DOI": _clean(r.get("doi")),
            "Summary": _summary_from_record(r),
        })
    return pd.DataFrame(rows)


def _to_excel_bytes(df: pd.DataFrame) -> bytes:
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Results")
        ws = writer.book["Results"]
        headers   = [c.value for c in ws[1]]
        link_cols = [i + 1 for i, h in enumerate(headers) if str(h).lower() in {"link", "linkedin"}]
        for col_idx in link_cols:
            for row_idx in range(2, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                if cell.value:
                    cell.hyperlink = str(cell.value)
                    cell.font = Font(color="0563C1", underline="single")

        # Wider columns for readability.
        for column_cells in ws.columns:
            header = str(column_cells[0].value or "")
            max_len = max(len(str(cell.value or "")) for cell in column_cells[:50])
            if header.lower() == "summary":
                ws.column_dimensions[column_cells[0].column_letter].width = 70
            else:
                ws.column_dimensions[column_cells[0].column_letter].width = min(max(max_len + 2, 12), 45)
    return out.getvalue()


def _to_pdf_bytes_vertical(df: pd.DataFrame, title: str) -> bytes:
    out  = io.BytesIO()
    doc  = SimpleDocTemplate(out, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story  = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for idx, row in df.iterrows():
        story.append(Paragraph(f"<b>Result {idx + 1}</b>", styles["Heading3"]))
        for col, val in row.items():
            text = _clean(val)
            if text:
                safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(f"<b>{col.title()}:</b> {safe}", styles["BodyText"]))
                story.append(Spacer(1, 4))
        story.append(Spacer(1, 10))
    doc.build(story)
    return out.getvalue()


def _to_word_bytes(df: pd.DataFrame, title: str) -> tuple[bytes, str]:
    if Document is not None:
        doc = Document()
        doc.add_heading(title, level=1)
        for idx, row in df.iterrows():
            doc.add_heading(f"Result {idx + 1}", level=2)
            for col, val in row.items():
                text = _clean(val)
                if text:
                    p = doc.add_paragraph()
                    p.add_run(f"{col.title()}: ").bold = True
                    p.add_run(text)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue(), ".docx"

    lines = [title, ""]
    for idx, row in df.iterrows():
        lines.append(f"Result {idx + 1}")
        for col, val in row.items():
            text = _clean(val)
            if text:
                lines.append(f"{col.title()}: {text}")
        lines.append("")
    return "\n".join(lines).encode("utf-8"), ".doc"


# ╔══════════════════════════════════════════════════════════╗
# ║  SESSION STATE INIT                                      ║
# ╚══════════════════════════════════════════════════════════╝

_MODE_DEFAULTS: dict[str, int] = {"Fast": 15, "Balanced": 25, "Deep": 40}

for _k, _v in [
    ("search_mode",     "Balanced"),
    ("coverage_slider", _MODE_DEFAULTS["Balanced"]),
    ("prompt_value",    ""),
    ("search_history",  []),
]:
    if _k not in st.session_state:
        st.session_state[_k] = _v


def _on_mode_change() -> None:
    """When the user changes Fast / Balanced / Deep, update coverage automatically."""
    selected_mode = st.session_state.search_mode
    st.session_state.coverage_slider = _MODE_DEFAULTS.get(selected_mode, 25)


def _set_prompt(text: str) -> None:
    st.session_state["prompt_value"] = text


# ╔══════════════════════════════════════════════════════════╗
# ║  SIDEBAR                                                 ║
# ╚══════════════════════════════════════════════════════════╝

with st.sidebar:
    st.markdown("## ⚙️ Search settings")

    # ── Mode (coverage auto-changes with mode) ───────────────────────────────
    st.radio(
        "Search mode",
        ["Fast", "Balanced", "Deep"],
        key="search_mode",
        horizontal=True,
        on_change=_on_mode_change,
        help=(
            "**Fast** — quickest, fewer sources.  "
            "**Balanced** — recommended for most searches.  "
            "**Deep** — best for niche or difficult topics."
        ),
    )
    mode = st.session_state.search_mode

    st.slider(
        "Search coverage",
        min_value=5,
        max_value=80,
        step=5,
        key="coverage_slider",
        help="How many candidate results to scan before filtering. Higher = wider net, slower search.",
    )
    search_coverage = int(st.session_state.coverage_slider)

    st.caption("Suggested coverage: Fast 15 · Balanced 25 · Deep 40")
    st.caption(f"Active: **{mode}** mode · coverage **{search_coverage}**")

    st.divider()

    # ── API Integrations ──────────────────────────────────────────────────────
    with st.expander("🔑 API integrations", expanded=False):

        # Step-by-step key guide with clickable provider links.
        st.markdown("""
<div class="key-guide">

<div class="kgp">
  <div class="kg-name">Groq <span class="kg-free">FREE TIER</span></div>
  <div class="kg-role">🧠 LLM reasoning — smarter query interpretation &amp; scoring</div>
  <ul class="kg-steps">
    <li><span class="sn">1.</span>Visit <a href="https://console.groq.com" target="_blank" rel="noopener noreferrer">console.groq.com</a> → sign up / log in</li>
    <li><span class="sn">2.</span>Left menu → <strong>API Keys</strong> → <strong>Create API key</strong></li>
    <li><span class="sn">3.</span>Copy key and paste below</li>
  </ul>
</div>

<div class="kgp">
  <div class="kg-name">Gemini <span class="kg-free">FREE TIER</span></div>
  <div class="kg-role">🧠 LLM reasoning — alternative to Groq</div>
  <ul class="kg-steps">
    <li><span class="sn">1.</span>Visit <a href="https://aistudio.google.com" target="_blank" rel="noopener noreferrer">aistudio.google.com</a> → sign in with Google</li>
    <li><span class="sn">2.</span>Click <strong>Get API key</strong> → <strong>Create API key</strong></li>
    <li><span class="sn">3.</span>Copy key and paste below</li>
  </ul>
</div>

<div class="kgp">
  <div class="kg-name">Exa</div>
  <div class="kg-role">🌐 Search engine — best for companies, papers &amp; LinkedIn</div>
  <ul class="kg-steps">
    <li><span class="sn">1.</span>Visit <a href="https://exa.ai" target="_blank" rel="noopener noreferrer">exa.ai</a> → sign up</li>
    <li><span class="sn">2.</span>Dashboard → <strong>API keys</strong> → <strong>Create key</strong></li>
    <li><span class="sn">3.</span>Free credits included on signup</li>
  </ul>
</div>

<div class="kgp">
  <div class="kg-name">Tavily</div>
  <div class="kg-role">🌐 Search engine — strong for news &amp; recent content</div>
  <ul class="kg-steps">
    <li><span class="sn">1.</span>Visit <a href="https://tavily.com" target="_blank" rel="noopener noreferrer">tavily.com</a> → sign up</li>
    <li><span class="sn">2.</span>Dashboard → <strong>API Keys</strong> → copy key</li>
    <li><span class="sn">3.</span>Free tier: 1,000 searches / month</li>
  </ul>
</div>

<div class="kgp">
  <div class="kg-name">SerpApi</div>
  <div class="kg-role">🌐 Search engine — Google results for broadest coverage</div>
  <ul class="kg-steps">
    <li><span class="sn">1.</span>Visit <a href="https://serpapi.com" target="_blank" rel="noopener noreferrer">serpapi.com</a> → sign up</li>
    <li><span class="sn">2.</span>Dashboard → copy <strong>Private API key</strong></li>
    <li><span class="sn">3.</span>Free plan: 100 searches / month</li>
  </ul>
</div>

</div>
""", unsafe_allow_html=True)

        groq_key    = st.text_input("Groq API key",   value="", type="password", placeholder="gsk_…")
        gemini_key  = st.text_input("Gemini API key", value="", type="password", placeholder="AIza…")
        exa_key     = st.text_input("Exa API key",    value="", type="password", placeholder="exa-…")
        tavily_key  = st.text_input("Tavily API key", value="", type="password", placeholder="tvly-…")
        serpapi_key = st.text_input("SerpApi key",    value="", type="password", placeholder="your key…")

        # Live provider status dots
        _key_map = {
            "DDG":     "always",
            "Groq":    groq_key    or _secret("GROQ_API_KEY"),
            "Gemini":  gemini_key  or _secret("GEMINI_API_KEY"),
            "Exa":     exa_key     or _secret("EXA_API_KEY"),
            "Tavily":  tavily_key  or _secret("TAVILY_API_KEY"),
            "SerpApi": serpapi_key or _secret("SERPAPI_KEY"),
        }
        dots_html = "".join(
            f'<span class="provider-dot {"on" if v.strip() else "off"}">'
            f'<span class="dot-circle {"on" if v.strip() else "off"}"></span>{name}'
            f'</span>'
            for name, v in _key_map.items()
        )
        st.markdown(f'<div class="provider-row">{dots_html}</div>', unsafe_allow_html=True)

        active = [n for n, v in _key_map.items() if v.strip() and n != "DDG"]
        if active:
            st.markdown(f"<div class='key-status'>✓ Active: {', '.join(active)} + DuckDuckGo</div>", unsafe_allow_html=True)
        else:
            st.caption("Running on DuckDuckGo only. Add keys above for wider coverage.")

    st.divider()

    with st.expander("⚙️ Advanced settings", expanded=False):
        requested_fields = st.multiselect(
            "Requested fields",
            ["website", "email", "phone", "linkedin", "summary", "presence_countries", "author"],
            default=["website", "email", "phone"],
            help="Which details to prioritize retrieving.",
        )
        min_confidence = st.slider(
            "Minimum relevance score",
            0, 100,
            25 if mode == "Fast" else 35,
            5,
            help="Higher = only strong matches retained. Lower = broader but noisier.",
        )
        uploaded_file   = st.file_uploader("Upload existing list for deduplication", type=["csv", "xlsx"])
        use_seed_dedupe = st.checkbox("Deduplicate against uploaded list", value=True)

    # Sidebar footer
    st.markdown(
        '<div class="sidebar-footer">'
        '🧭 Research Navigator<br>'
        'Tip: <strong>Balanced</strong> is best for most searches.<br>'
        'Add Groq + Exa for strongest results.'
        '</div>',
        unsafe_allow_html=True,
    )


# ╔══════════════════════════════════════════════════════════╗
# ║  MAIN PAGE                                               ║
# ╚══════════════════════════════════════════════════════════╝

# ── Hero ──────────────────────────────────────────────────────────────────────
st.markdown(
    '<div class="hero-wrap">'
    '<div>'
    '<div class="hero-title">🧭 Research Navigator</div>'
    '<div class="hero-subtitle">Find companies, papers, LinkedIn profiles &amp; more — describe your search in English or Arabic.</div>'
    '</div>'
    '<div class="hero-badge">⚡ AI-powered search agent</div>'
    '</div>',
    unsafe_allow_html=True,
)

# ── 3-tier capability card ────────────────────────────────────────────────────
st.markdown("""
<div class="tier-card">
  <div class="tier-card-header">
    <div>
      <div class="tier-card-title">Unlock stronger results with API integrations</div>
      <div class="tier-card-sub">Works out of the box — adding free API keys significantly improves coverage and accuracy.</div>
    </div>
    <div class="sidebar-cta">☰ &nbsp;Sidebar → API integrations</div>
  </div>
  <div class="tiers-grid">
    <div class="tier-cell tier-free">
      <div class="tier-cell-label">Free · No keys needed</div>
      <div class="tier-cell-name">DuckDuckGo</div>
      <div class="tier-cell-desc">Basic web search. Good starting point for common or broad queries.</div>
    </div>
    <div class="tier-cell tier-llm">
      <div class="tier-cell-label">+ LLM reasoning (free)</div>
      <div class="tier-cell-name">Groq · Gemini</div>
      <div class="tier-cell-desc">Smarter interpretation, better filtering &amp; relevance scoring. Both offer free tiers.</div>
    </div>
    <div class="tier-cell tier-full">
      <div class="tier-cell-label">+ Search engines</div>
      <div class="tier-cell-name">Exa · Tavily · SerpApi</div>
      <div class="tier-cell-desc">Wider coverage, more sources, much stronger results for companies &amp; papers.</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Quick-fill example chips ──────────────────────────────────────────────────
_EXAMPLES: list[tuple[str, str]] = [
    ("🏭 O&G companies",       "Find oil and gas service companies in Egypt with website and email."),
    ("📄 ESP papers",           "Find academic papers about electrical submersible pumps with authors and abstract."),
    ("👤 LinkedIn engineers",   "Find LinkedIn profiles of petroleum engineers in Saudi Arabia."),
    ("🌍 Digital O&G vendors",  "Find digital software companies serving the oil and gas industry in the Middle East, excluding USA-headquartered firms."),
]
st.markdown('<div class="chips-label">Quick examples — click to fill</div>', unsafe_allow_html=True)
_ecols = st.columns(len(_EXAMPLES))
for _col, (_label, _prompt) in zip(_ecols, _EXAMPLES):
    if _col.button(_label, use_container_width=True, key=f"ex_{_label}"):
        _set_prompt(_prompt)
        st.rerun()

# ── Search box ────────────────────────────────────────────────────────────────
st.markdown('<div class="search-shell">', unsafe_allow_html=True)
prompt = st.text_area(
    "What would you like to research?",
    value=st.session_state.get("prompt_value", ""),
    height=210,
    placeholder=(
        "Describe what you want to find in English or Arabic.\n\n"
        "Examples:\n"
        "• Find service companies in upstream oil and gas in Egypt — website, email, phone.\n"
        "• Find SPE papers about ROP prediction for drill bits with authors and DOI.\n"
        "• ابحث عن شركات خدمات البترول في مصر مع الموقع الإلكتروني والإيميل."
    ),
)
run_btn = st.button("🔍  Start search", type="primary", use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

# ── Search history ────────────────────────────────────────────────────────────
_history: list[str] = st.session_state.get("search_history", [])
if _history:
    st.markdown('<div class="history-label">🕐 Recent searches</div>', unsafe_allow_html=True)
    _hcols = st.columns(min(len(_history), 5))
    for _idx, (_hcol, _hprompt) in enumerate(zip(_hcols, _history[:5])):
        _short = _hprompt[:42] + ("…" if len(_hprompt) > 42 else "")
        if _hcol.button(_short, key=f"hist_{_idx}", use_container_width=True):
            _set_prompt(_hprompt)
            st.rerun()


# ╔══════════════════════════════════════════════════════════╗
# ║  SEARCH EXECUTION                                        ║
# ╚══════════════════════════════════════════════════════════╝

if run_btn:
    st.session_state["prompt_value"] = prompt
    if not prompt.strip():
        st.warning("Please enter a search query above.")
    else:
        # Save to history (deduplicate, most-recent first, cap at 5)
        _hist = [h for h in st.session_state["search_history"] if h != prompt]
        st.session_state["search_history"] = [prompt] + _hist[:4]

        user_keys = {
            "groq_api_key":   groq_key    or _secret("GROQ_API_KEY"),
            "gemini_api_key": gemini_key  or _secret("GEMINI_API_KEY"),
            "exa_api_key":    exa_key     or _secret("EXA_API_KEY"),
            "tavily_api_key": tavily_key  or _secret("TAVILY_API_KEY"),
            "serpapi_key":    serpapi_key or _secret("SERPAPI_KEY"),
        }

        llm_client = FreeLLMClient(
            groq_api_key=user_keys["groq_api_key"],
            gemini_api_key=user_keys["gemini_api_key"],
        )

        task_spec = parse_task_prompt_llm_first(prompt, llm=llm_client)
        task_spec.mode        = mode
        task_spec.max_results = int(search_coverage)
        if requested_fields:
            task_spec.target_attributes = list(dict.fromkeys(requested_fields))
        elif not task_spec.target_attributes:
            task_spec.target_attributes = ["website"]

        provider_settings = ProviderSettings(
            use_ddg=True,
            use_exa=bool(user_keys["exa_api_key"]),
            use_tavily=bool(user_keys["tavily_api_key"]),
            use_serpapi=bool(user_keys["serpapi_key"]),
            use_firecrawl=False,
            use_llm_parser=bool(llm_client.available_backends()),
            use_uploaded_seed_dedupe=use_seed_dedupe,
        )
        uploaded_df = _read_file(uploaded_file)

        # ── Named progress stages ─────────────────────────────────────────────
        _STAGES = [
            "🧠 Interpreting your query…",
            "🌐 Querying search engines…",
            "🌐 Expanding coverage…",
            "🔍 Filtering by relevance…",
            "✅ Verifying matches with AI…",
            "📦 Packaging results…",
        ]
        _stage_idx: dict[str, int] = {"i": 0}
        _progress_container = st.empty()
        _progress_bar       = _progress_container.progress(0, text=_STAGES[0])

        def _progress(msg: str) -> None:
            _stage_idx["i"] = min(_stage_idx["i"] + 1, len(_STAGES) - 1)
            pct = int((_stage_idx["i"] / len(_STAGES)) * 90)
            _progress_bar.progress(pct, text=_STAGES[_stage_idx["i"]])

        result = SearchOrchestrator().run_task(
            task_spec=task_spec,
            provider_settings=provider_settings,
            uploaded_df=uploaded_df,
            budget_overrides={},
            min_confidence_score=int(min_confidence),
            user_keys=user_keys,
            progress_callback=_progress,
        )
        _progress_bar.progress(100, text="✅ Search complete!")
        _progress_container.empty()

        st.session_state["last_result"]   = result
        st.session_state["last_mode"]     = mode
        st.session_state["last_coverage"] = search_coverage


# ╔══════════════════════════════════════════════════════════╗
# ║  RESULTS                                                 ║
# ╚══════════════════════════════════════════════════════════╝

result = st.session_state.get("last_result")
if result:
    task_meta = result.get("task_spec", {}) or {}
    records   = result.get("records",   []) or []

    _render_search_summary(task_meta)
    _render_results_bar(
        n_results=len(records),
        mode=st.session_state.get("last_mode", mode),
        coverage=st.session_state.get("last_coverage", search_coverage),
        task_type=task_meta.get("task_type", ""),
    )

    if records:
        display_df = _build_display_df(records, task_meta)
        export_df  = _build_export_df(records, task_meta)

        is_papers = (
            task_meta.get("task_type") == "document_research"
            or (task_meta.get("target_entity_types") or [""])[0] == "paper"
        )
        if is_papers:
            results_tab, summaries_tab = st.tabs(["📊 Results", "📄 Paper summaries"])
        else:
            results_tab   = st.container()
            summaries_tab = None

        with results_tab:
            column_config: dict = {}
            if "Link" in display_df.columns:
                column_config["Link"]     = st.column_config.LinkColumn("Link")
            if "LinkedIn" in display_df.columns:
                column_config["LinkedIn"] = st.column_config.LinkColumn("LinkedIn")
            st.dataframe(display_df, use_container_width=True, hide_index=True, column_config=column_config)

            st.markdown('<div class="download-box">', unsafe_allow_html=True)
            st.markdown("### 📥 Download results")
            st.caption("Exports include: name, link, email, phone, LinkedIn, and summary.")
            c1, c2, c3 = st.columns(3)
            _label       = _human_task(task_meta.get("task_type", ""))
            excel_bytes  = _to_excel_bytes(export_df)
            pdf_bytes    = _to_pdf_bytes_vertical(export_df, f"Research Navigator — {_label}")
            word_bytes, word_ext = _to_word_bytes(export_df, f"Research Navigator — {_label}")
            with c1:
                st.download_button(
                    "📊 Excel",
                    data=excel_bytes,
                    file_name=_normalize_filename("results", ".xlsx"),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    on_click="ignore",
                )
            with c2:
                st.download_button(
                    "📝 Word",
                    data=word_bytes,
                    file_name=_normalize_filename("results", word_ext),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    on_click="ignore",
                )
            with c3:
                st.download_button(
                    "📄 PDF",
                    data=pdf_bytes,
                    file_name=_normalize_filename("results", ".pdf"),
                    mime="application/pdf",
                    use_container_width=True,
                    on_click="ignore",
                )
            st.markdown('</div>', unsafe_allow_html=True)

        if summaries_tab is not None:
            with summaries_tab:
                summary_df = _build_paper_summaries_df(records)

                st.markdown("### 📄 Paper summaries")
                st.caption("Readable summaries with title, authors, year, DOI, source link, and summary text.")

                for _, row in summary_df.iterrows():
                    title = _clean(row.get("Title")) or "Untitled"
                    link = _clean(row.get("Link"))
                    authors = _clean(row.get("Authors"))
                    year = _clean(row.get("Year"))
                    doi = _clean(row.get("DOI"))
                    summary = _clean(row.get("Summary"))

                    st.markdown(f"#### {title}")

                    meta_parts = []
                    if authors:
                        meta_parts.append(f"**Authors:** {authors}")
                    if year:
                        meta_parts.append(f"**Year:** {year}")
                    if doi:
                        meta_parts.append(f"**DOI:** {doi}")

                    if meta_parts:
                        st.markdown(" | ".join(meta_parts))

                    if link:
                        st.markdown(f"[🔗 Open paper/source]({link})")

                    st.write(summary or "No summary available.")
                    st.divider()

                st.markdown('<div class="download-box">', unsafe_allow_html=True)
                st.markdown("### 📥 Download paper summaries")
                st.caption("Exports are formatted for reading, not only spreadsheet viewing.")

                s1, s2, s3 = st.columns(3)

                summary_excel_bytes = _to_excel_bytes(summary_df)
                summary_pdf_bytes = _to_pdf_bytes_vertical(
                    summary_df,
                    "Research Navigator — Paper Summaries",
                )
                summary_word_bytes, summary_word_ext = _to_word_bytes(
                    summary_df,
                    "Research Navigator — Paper Summaries",
                )

                with s1:
                    st.download_button(
                        "📊 Excel summaries",
                        data=summary_excel_bytes,
                        file_name=_normalize_filename("paper_summaries", ".xlsx"),
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        on_click="ignore",
                    )

                with s2:
                    st.download_button(
                        "📝 Word summaries",
                        data=summary_word_bytes,
                        file_name=_normalize_filename("paper_summaries", summary_word_ext),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        on_click="ignore",
                    )

                with s3:
                    st.download_button(
                        "📄 PDF summaries",
                        data=summary_pdf_bytes,
                        file_name=_normalize_filename("paper_summaries", ".pdf"),
                        mime="application/pdf",
                        use_container_width=True,
                        on_click="ignore",
                    )

                st.markdown('</div>', unsafe_allow_html=True)
    else:
        _render_empty_state(task_meta.get("task_type", ""))
